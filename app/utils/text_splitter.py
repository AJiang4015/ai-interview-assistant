import re
from pathlib import Path
from typing import Optional

from app.utils.logger import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".md", ".pdf", ".docx"}


class DocumentParser:
    """Parse different document formats into plain text."""

    @staticmethod
    def parse_file(file_path: Path) -> Optional[str]:
        ext = file_path.suffix.lower()
        try:
            if ext == ".md":
                return DocumentParser._parse_markdown(file_path)
            elif ext == ".pdf":
                return DocumentParser._parse_pdf(file_path)
            elif ext == ".docx":
                return DocumentParser._parse_docx(file_path)
            else:
                return None
        except Exception as e:
            logger.error(f"Failed to parse {file_path.name}: {type(e).__name__}: {e}")
            return None

    @staticmethod
    def _parse_markdown(file_path: Path) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    @staticmethod
    def _parse_pdf(file_path: Path) -> str:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())
        return "\n\n".join(text_parts)

    @staticmethod
    def _parse_docx(file_path: Path) -> str:
        from docx import Document
        doc = Document(str(file_path))
        text_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)


class MarkdownSplitter:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_file(self, file_path: str | Path) -> list[dict]:
        file_path = Path(file_path)
        text = DocumentParser.parse_file(file_path)
        if text is None:
            return []
        return self.split_text(text, source_file=file_path.name)

    def split_text(self, text: str, source_file: str) -> list[dict]:
        sections = self._split_by_headers(text)
        chunks = []
        chunk_index = 0
        for section_title, section_content in sections:
            blocks = self._split_into_blocks(section_title, section_content)
            for block in blocks:
                chunks.append({
                    "content": block,
                    "source_file": source_file,
                    "chunk_index": chunk_index
                })
                chunk_index += 1
        return chunks

    def _split_by_headers(self, text: str) -> list[tuple[str, str]]:
        pattern = r'^(#{1,3})\s+(.+)$'
        lines = text.split("\n")
        sections = []
        current_title = "Preamble"
        current_lines = []
        for line in lines:
            m = re.match(pattern, line)
            if m:
                if current_lines:
                    sections.append((current_title, "\n".join(current_lines)))
                current_title = m.group(2).strip()
                current_lines = []
            else:
                current_lines.append(line)
        if current_lines:
            sections.append((current_title, "\n".join(current_lines)))
        if not sections:
            sections.append(("Full Text", text))
        return sections

    def _split_into_blocks(self, title: str, content: str) -> list[str]:
        header = f"## {title}\n\n"
        full_text = header + content.strip()
        blocks = []
        if len(full_text) <= self.chunk_size:
            blocks.append(full_text)
            return blocks

        paragraphs = re.split(r'\n\s*\n', full_text)
        current = ""
        for para in paragraphs:
            if len(current) + len(para) + 2 <= self.chunk_size:
                current = current + "\n\n" + para if current else para
            else:
                if current:
                    blocks.append(current)
                if len(para) > self.chunk_size:
                    for i in range(0, len(para), self.chunk_size - self.chunk_overlap):
                        blocks.append(para[i:i + self.chunk_size])
                else:
                    current = para
        if current:
            blocks.append(current)
        return blocks

    def scan_md_files(self, directory: str | Path) -> list[Path]:
        directory = Path(directory)
        if not directory.exists():
            return []
        files = []
        for ext in SUPPORTED_EXTENSIONS:
            for f in directory.glob(f"*{ext}"):
                if not f.name.startswith("~$"):
                    files.append(f)
        return sorted(files)
